#!/usr/bin/env python3
"""Cross-document priority consistency check.

The data files are the source of truth. This verifies that every other place a
milestone appears agrees with them: the requirements document's priority column
and its priority-summary appendix, the epic map's milestone labels, its
per-milestone requirement groupings and its requirement index, the HTML board,
and any prose in either document that names a requirement and a milestone in
the same breath.

Usage: consistency-check.py <dir> <fr.docx> <epic.docx> <board.html>
"""
import sys, re, json, os, subprocess, collections
import docx

DIR, FRDOC, EPICDOC, BOARD = sys.argv[1:5]

# ---------------------------------------------------------------- truth
truth = json.loads(subprocess.check_output(["node", "-e", f"""
const s=[...require('{DIR}/reqs-part1.js'),...require('{DIR}/reqs-part2.js')];
const A=require('{DIR}/appendices.js');
const E=require('{DIR}/epics.js');
const pri={{}}, notes={{}}, feat={{}};
s.forEach(x=>x.reqs.forEach(r=>{{pri[r[0]]=r[4]; notes[r[0]]=r[3]||'';}}));
E.forEach(e=>e.features.forEach(f=>{{f[2].forEach(i=>feat[i]=f[0]);}}));
const bl={{}};   // the backlog was emptied at Rev 1.18; kept so the shape is stable
console.log(JSON.stringify({{pri,notes,feat,bl,
  epics:E.map(e=>({{code:e.code,features:e.features.map(f=>({{id:f[0],fr:f[2],
    intent:f[4],ux:f[5],split:f[6]||null}}))}}))}}));
"""]).decode())
PRI, NOTES, FEAT, BL = truth["pri"], truth["notes"], truth["feat"], truth["bl"]
MS = ["1", "2", "3", "4", "5"]
fails, checks = [], 0


def check(cond, msg):
    global checks
    checks += 1
    if not cond:
        fails.append(msg)


def label(ids):
    n = sorted({int(PRI[i]) for i in ids if PRI[i] in MS})
    tbd = any(PRI[i] == "TBD" for i in ids)
    if not n:
        return "TBD" if tbd else "—"
    s = "M%d" % n[0] + ("–M%d" % n[-1] if n[-1] != n[0] else "")
    return s + (" +TBD" if tbd else "")


# ---------------------------------------------------------------- FR doc
d = docx.Document(FRDOC)
fr_pri, fr_feat, apxJ = {}, {}, {}
for t in d.tables:
    h = [c.text.strip() for c in t.rows[0].cells]
    if not h:
        continue
    if h[0] == "Req. ID" and "Feat." in h:
        pi, fi = h.index("Pri."), h.index("Feat.")
        for r in t.rows[1:]:
            c = [x.text.strip() for x in r.cells]
            fr_pri[c[0]] = c[pi]
            fr_feat[c[0]] = c[fi]
    if h and h[0] == "Section" and any("1" == x for x in h):
        for r in t.rows[1:]:
            c = [x.text.strip() for x in r.cells]
            apxJ[c[0]] = c

check(set(fr_pri) == set(PRI), "requirements doc holds a different set of identifiers from the data")
for k, v in fr_pri.items():
    check(v == PRI.get(k), f"requirements doc shows {k} at {v}, data says {PRI.get(k)}")
for k, v in fr_feat.items():
    check(v == FEAT.get(k), f"requirements doc shows {k} in feature {v}, epic map says {FEAT.get(k)}")

dist = collections.Counter(PRI.values())
apxJ_total = None
for k, c in apxJ.items():
    if k.lower().startswith("all") or k.lower().startswith("total"):
        apxJ_total = c
if apxJ_total:
    nums = [x for x in apxJ_total[1:] if x.isdigit()]
    check(str(len(PRI)) in nums or str(sum(dist.values())) in nums,
          f"priority summary total {nums} does not match {len(PRI)} requirements")

# ---------------------------------------------------------------- epic doc
e = docx.Document(EPICDOC)
ep_feat_ms, ep_groups, ep_index = {}, {}, {}
for t in e.tables:
    h = [c.text.strip() for c in t.rows[0].cells]
    if not h:
        continue
    if h[0] == "Ref." and "Feature and intent" in h:
        mi, ri = h.index("M"), h.index("Requirements")
        for r in t.rows[1:]:
            c = [x.text.strip() for x in r.cells]
            if not re.match(r"^F\d\d\.\d+$", c[0]):
                continue
            ep_feat_ms[c[0]] = c[mi].replace("\n", "").replace("- ", "").replace("-", "–")
            ep_groups[c[0]] = c[ri]
    if h[0] == "Req. ID" and "Feature name" in h:
        for r in t.rows[1:]:
            c = [x.text.strip() for x in r.cells]
            ep_index[c[0]] = (c[1], c[3])

for ep in truth["epics"]:
    for f in ep["features"]:
        want = label(f["fr"])
        got = ep_feat_ms.get(f["id"])
        check(got is not None, f"epic map has no row for {f['id']}")
        # The milestone cell may render "+TBD" on its own line, so whitespace is
        # not significant when comparing the label to the derived one.
        norm = lambda x: x.replace("–", "-").replace("—", "-").replace(" ", "")
        if got is not None:
            check(norm(got) == norm(want),
                  f"epic map shows {f['id']} as {got!r}, derived from the data it is {want!r}")
        # per-milestone requirement grouping
        cellt = ep_groups.get(f["id"], "")
        by = collections.defaultdict(list)
        for i in f["fr"]:
            by[PRI[i]].append(i.replace("FR-", ""))
        for m, lst in by.items():
            tag = "TBD" if m == "TBD" else "M" + m
            for rid in lst:
                check(rid in cellt, f"{f['id']} requirements cell omits {rid}")
            check(tag in cellt, f"{f['id']} requirements cell has no {tag} group")

for k, (m, feat) in ep_index.items():
    check(m == PRI.get(k), f"epic map index shows {k} at {m}, data says {PRI.get(k)}")
    check(feat == FEAT.get(k), f"epic map index shows {k} in {feat}, data says {FEAT.get(k)}")

# ---------------------------------------------------------------- board
board = open(BOARD, encoding="utf8").read()
m = re.search(r"const D=(\{.*?\});\nconst MSNAME", board, re.S)
D = json.loads(m.group(1))
for k, v in D["req"].items():
    check(v["m"] == PRI.get(k), f"board shows {k} at {v['m']}, data says {PRI.get(k)}")
check(len(D["req"]) == len(PRI), "board carries a different number of requirements")

# ------------------------------------------------- appendix cross-references
# Every "Appendix X" the document says must be an appendix the document has. Rev 1.17 and
# every revision before it carried FR-LCM-018 (then BL-15) citing "Appendix H", which has
# never existed: the headings ran A to G then skipped to I. A reader sizing that
# requirement would have looked for an appendix that was not there.
_frdoc = docx.Document(FRDOC)
_alltext = "\n".join([p.text for p in _frdoc.paragraphs] +
                     [c.text for t in _frdoc.tables for r in t.rows for c in r.cells])
have = set(re.findall(r"^Appendix ([A-Z]) [–-]", _alltext, re.M))
check(bool(have), "the FR document has no appendix headings")
cited = set(re.findall(r"Appendix ([A-Z])\b", _alltext))
for letter in sorted(cited - have):
    check(False, f"the document cites Appendix {letter}, which it does not contain "
                 f"(it has {', '.join(sorted(have))})")


# ------------------------------------------------- appendix subsection labels
# Appendix E's three subsections shipped labelled F.1 to F.3 and Appendix F's labelled G.1 to
# G.3 — each carrying the *next* appendix's letter, left behind when the appendices were
# re-lettered. Six wrong labels across two shipped appendices of a submission document, and
# no guard looked: the cross-reference check above asks whether a cited appendix exists, never
# whether a subsection sits under the appendix it names.
_cur = None
for _line in _alltext.splitlines():
    _m = re.match(r"^Appendix ([A-Z]) [\u2013-]", _line)
    if _m:
        _cur = _m.group(1)
        continue
    _s = re.match(r"^([A-Z])\.(\d+)\s+\S", _line)
    # only an appendix letter counts. The review register at the front numbers its groups
    # R.1 to R.6, which is a deliberate prefix and not an appendix that has drifted.
    if _s and _s.group(1) not in have:
        continue
    if _s and _cur and _s.group(1) != _cur:
        check(False, f"Appendix {_cur} contains a subsection labelled "
                     f"{_s.group(1)}.{_s.group(2)} — {_line.strip()[:60]!r}")
    elif _s and _cur:
        check(True, f"appendix subsection {_s.group(1)}.{_s.group(2)} matches its appendix")
# ------------------------------------------------- the register is as long as the data
# The review register's summary states how many open items each group holds, and those
# numbers are computed from review.js at build time. So they are a fingerprint of the data
# the document was built from: if the data has moved on, they disagree.
#
# Rev 1.18 shipped saying 47 open items while review.js held 48. Q-100 had been added to the
# data after the document was built, and the document was never rebuilt — so the deliverable
# and the source that defines it disagreed, which is the one thing section 1 of CLAUDE.md
# exists to prevent. Every other guard passed, because they all compare a milestone or an
# identifier and none of them counts.
_rev = json.loads(subprocess.check_output(
    ["node", "-e", f"console.log(JSON.stringify(require('{DIR}/review.js')))"]).decode())
_tbd = sum(1 for v in PRI.values() if v == "TBD")
_groups = [("R.1 Priorities not assigned", _tbd),
           ("R.2 Parameter values not set", len(_rev["parameters"])),
           ("R.3 Conflicts between documents", len(_rev["conflicts"])),
           ("R.4 Scope questions", len(_rev["scope"])),
           ("R.5 Risk Analysis edits", len(_rev["raEdits"])),
           ("R.6 Consequences of deferral", len(_rev["consequences"]))]
_summary = {}
for t in _frdoc.tables:
    for r in t.rows:
        c = [x.text.strip() for x in r.cells]
        if len(c) >= 2 and (c[0].startswith("R.") or c[0] == "All groups"):
            _summary[c[0]] = c[1]
check(bool(_summary), "the FR document has no review-register summary to check")
for label, n in _groups:
    check(_summary.get(label) == str(n),
          f"review register says {label} holds {_summary.get(label)!r}, the data holds {n} "
          f"— the document was built from data that has since changed; rebuild")
check(_summary.get("All groups") == str(sum(n for _, n in _groups)),
      f"review register total is {_summary.get('All groups')!r}, the data totals "
      f"{sum(n for _, n in _groups)} — rebuild")

# ------------------------------------------------- no dead identifier class
# BL-nn ceased to exist at Rev 1.18, when the backlog was dissolved into milestone 5. A
# reader who meets "deferred (BL-24)" in a note has no way to resolve it: there is no
# backlog appendix and no such identifier. Eleven references survived the remap in prose
# the remap did not visit — requirement notes, epic-map design notes, the review register
# and the configuration register — and every one of them shipped in Rev 1.18 with 1,735
# checks passing, because nothing looked at the built documents for a retired class.
#
# Checked on the artefacts rather than the sources, so it cannot be evaded by a path
# nobody thought of.
# Two kinds of mention are legitimate and everything else is stale. "formerly BL-nn" is the
# audit trail every restored requirement carries, and id-guard.js depends on it. Beyond that
# there is an explicit allowlist, not a heuristic: a mention that is genuinely historical has
# to be recorded here with its reason, the same discipline as id-manifest.json.
BL_ALLOWED = {
    "BL-38": "withdrawn as a backlog item and never a requirement, so Appendix D lists it "
             "under the only identifier it ever had",
    "BL-13": "removed from the backlog at Rev 1.7, which is why it is a gap in the "
             "numbering; the decision log records that history and must keep the number",
}
# ------------------------------------------------- no row is formerly itself
# The audit trail says where an identifier came from, so a row claiming it was formerly its
# own identifier carries no information and destroys what was there. The Rev 1.18 migration
# did this to 41 requirement rows with a blind substitution; that was caught and repaired,
# but the same substitution had also hit the review register's reference column, where five
# rows survived saying "FR-CFG-003, formerly FR-CFG-003" and shipped.
for m in re.finditer(r"(FR-[A-Z]{3}-\d{3})[^\n]{0,40}?formerly \1\b", _alltext):
    check(False, f"a row says it was formerly itself: {m.group(0)!r} — the marker should name "
                 f"the identifier it came from, or be dropped")

# The word itself, not just the identifiers. Rev 1.18 shipped fifteen mentions of a backlog
# that no longer exists — a requirement note citing "the backlog appendix" that had been
# removed, a configuration register asserting the blocked state had no defined response when
# it had just come into scope, and a dozen "deferred to the backlog" claims that should read
# milestone 5. Only genuine history may say the word, and it has to be allowlisted here.
BACKLOG_ALLOWED = [
    "Was a deferred backlog item and is not carried into milestone 5",  # Appendix D, BL-38
    "Are backlog and requirement references stable?",                   # decision log, Q-62
]
for m in re.finditer(r"backlog", _alltext, re.I):
    window = _alltext[max(0, m.start() - 120):m.end() + 120]
    check(any(a in window for a in BACKLOG_ALLOWED),
          "the FR document still refers to a backlog, which ceased to exist at Rev 1.18: "
          f"…{_alltext[max(0, m.start() - 70):m.end() + 70]}…".replace("\n", " "))

_epicdoc = docx.Document(EPICDOC)
_epictext = "\n".join([p.text for p in _epicdoc.paragraphs] +
                      [c.text for t in _epicdoc.tables for r in t.rows for c in r.cells])
for label, body in (("the FR document", _alltext),
                    ("the epic map", _epictext),
                    ("the board", open(BOARD, encoding="utf-8").read())):
    dead = sorted({m.group(0) for m in re.finditer(r"(?<!formerly )\bBL-\d{1,3}\b", body)}
                  - set(BL_ALLOWED))
    for d in dead:
        check(False, f"{label} names {d} outside a 'formerly' marker, but BL-nn ceased to "
                     f"exist at Rev 1.18 — every backlog item now carries an FR identifier")

# ------------------------------------------------- the prefix legend
# Section 2 of the FR document carries a table of identifier prefixes and functional
# areas. Its wording is authored, but the *set* of prefixes must equal the set of live
# sections exactly. Rev 1.16 shipped listing FR-SUR — Post-test survey, a section that
# does not exist: all five FR-SUR identifiers were withdrawn or superseded to BL-29..32.
live_prefixes = {rid.split("-")[1] for rid in PRI}
_frdoc = docx.Document(FRDOC)
legend = set()
for t in _frdoc.tables:
    rows = [[c.text.strip() for c in r.cells] for r in t.rows]
    if rows and rows[0][:2] == ["Prefix", "Functional area"]:
        legend = {r[0].split("-")[1] for r in rows[1:] if r[0].startswith("FR-")}
check(bool(legend), "the FR document has no prefix legend table")
for extra in sorted(legend - live_prefixes):
    check(False, f"prefix legend lists FR-{extra}, but no section uses that prefix")
for gone in sorted(live_prefixes - legend):
    check(False, f"section prefix FR-{gone} is missing from the prefix legend")

# ------------------------------------------------- document revisions
# version.js is the only place a revision may be declared. Every self-description in
# the built documents must agree with it, and no builder may hold a literal.
VER = json.loads(subprocess.check_output(
    ["node", "-e", f"console.log(JSON.stringify(require('{DIR}/version.js')))"]).decode())

def selfdesc(path):
    d = docx.Document(path)
    rows, hf = {}, []
    for t in d.tables:
        for r in t.rows:
            c = [x.text.strip() for x in r.cells]
            if c and c[0] in ("Revision", "Derived from"):
                rows[c[0]] = c[1]
    for sec in d.sections:
        for para in sec.header.paragraphs + sec.footer.paragraphs:
            if para.text.strip():
                hf.append(para.text.strip())
    return rows, " | ".join(hf)

fr_rows, fr_hf = selfdesc(FRDOC)
check(fr_rows.get("Revision", "").startswith(VER["FR"] + " "),
      f'FR front matter says revision {fr_rows.get("Revision")!r}, version.js says {VER["FR"]}')
check(f'Rev {VER["FR"]}' in fr_hf,
      f'FR header/footer does not carry Rev {VER["FR"]}: {fr_hf!r}')

ep_rows, ep_hf = selfdesc(EPICDOC)
check(ep_rows.get("Revision", "").startswith(VER["EPIC"] + " "),
      f'epic map front matter says revision {ep_rows.get("Revision")!r}, version.js says {VER["EPIC"]}')
check(f'Rev {VER["EPIC"]}' in ep_hf,
      f'epic map header/footer does not carry Rev {VER["EPIC"]}: {ep_hf!r}')
check(ep_rows.get("Derived from", "").endswith("Rev " + VER["FR"]),
      f'epic map "Derived from" says {ep_rows.get("Derived from")!r}, but the FR revision is {VER["FR"]}')

board_html = open(BOARD, encoding="utf-8").read()
check(f'Rev {VER["EPIC"]}' in board_html, f'board does not carry epic Rev {VER["EPIC"]}')
check(f'Rev {VER["FR"]}' in board_html, f'board does not carry FR Rev {VER["FR"]}')

# The cover page states the document's own revision in a bare paragraph, outside any
# table, so the front-matter check above cannot see it. Rev 1.16 shipped with a cover
# reading "Revision 1.15", and the epic map Rev 1.11 with "Revision 1.9": both builders
# held a frozen literal on the line directly above an interpolated one. This is the
# mistake version.js was created to prevent, so the cover is now checked on its own.
COVER = re.compile(r"Document no\.\s+QACR-APP-(?:FR|EPIC)-01\D{0,20}Revision\s+(\d+\.\d+)")


def cover_revision(path):
    for para in docx.Document(path).paragraphs:
        m = COVER.search(para.text)
        if m:
            return m.group(1)
    return None


for _path, _label, _num in ((FRDOC, "FR", VER["FR"]), (EPICDOC, "epic map", VER["EPIC"])):
    _rev = cover_revision(_path)
    check(_rev is not None, f"{_label} has no cover-page revision line for this guard to check")
    if _rev is not None:
        check(_rev == _num, f"{_label} cover page says Revision {_rev}, version.js says {_num}")

# no builder may contain a literal revision of our own documents.
# "Rev(ision)?" because the frozen cover literals said "Revision 1.15", which an
# earlier "Rev\s*1\.\d+" did not match: after "Rev" came "ision", not a digit.
OWN = re.compile(r'(QACR-APP-(?:FR|EPIC)-01[^"`\n]{0,60}Rev(?:ision)?\s*1\.\d+)|("1\.\d+ — )')
for b in ("build.js", "build-epics.js", "build-board.js"):
    src = open(os.path.join(DIR, b), encoding="utf-8").read()
    for m in OWN.finditer(src):
        check(False, f"{b} holds a literal revision {m.group(0)!r} — it must come from version.js")

# ------------------------------------------------- prose claims about milestones
# Anything of the form "<ID> ... milestone N" or "milestone N ... <ID>" within one
# sentence, in any note, intent, split line or register row.
prose = []
for k, v in NOTES.items():
    prose.append((f"note on {k}", v))
for ep in truth["epics"]:
    for f in ep["features"]:
        prose.append((f["id"] + " intent", f["intent"]))
        prose.append((f["id"] + " design note", f["ux"]))
        for lab, txt in (f["split"] or []):
            prose.append((f"{f['id']} split {lab}", txt))

# A sentence is only checkable if it names exactly one requirement and exactly one
# milestone; anything more tangled is reported for a human to read rather than
# guessed at, because attributing the number to the right identifier is ambiguous.
IDPAT = re.compile(r"FR-[A-Z]{3}-\d{3}")
MSPAT = re.compile(r"milestone\s+(\d)", re.I)
claims, unverifiable = 0, []
for where, txt in prose:
    for sent in re.split(r"(?<=[.;])\s+", txt or ""):
        ids = IDPAT.findall(sent)
        mss = MSPAT.findall(sent)
        if not ids or not mss:
            continue
        if len(set(ids)) == 1 and len(set(mss)) == 1:
            claims += 1
            check(PRI.get(ids[0]) == mss[0],
                  f"{where} says {ids[0]} is milestone {mss[0]}, data says {PRI.get(ids[0])}")
        else:
            unverifiable.append(f"{where}: “{sent.strip()[:150]}”")

# ---------------------------------------------------------------- report
print(f"checks run: {checks}   prose milestone claims verified: {claims}")
if unverifiable:
    print(f"\nprose naming requirements and milestones together but too tangled to check ({len(unverifiable)}) —")
    for u in unverifiable:
        print("   ", u)
print(f"requirements: {len(PRI)}   future development: {sum(1 for v in PRI.values() if v == '5')}   distribution: {dict(sorted(dist.items()))}")
if unverifiable:
    print(f"\nsentences naming several requirements or milestones — read these by eye ({len(unverifiable)}):")
    for u in unverifiable:
        print("  " + u)
if fails:
    print(f"\nDISCREPANCIES ({len(fails)}):")
    for f in fails:
        print("  " + f)
    sys.exit(1)
print("\nno discrepancies: every milestone in both documents and the board agrees with the data")
