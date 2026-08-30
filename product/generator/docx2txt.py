#!/usr/bin/env python3
"""Renders a .docx as plain text so `git diff` can show what changed inside it.

Wired up as a textconv diff driver, so `git diff` and `git log -p` on a Word file
show readable document changes instead of "Binary files differ". See .gitattributes.

Enable once per clone:

    git config diff.docx.textconv "python3 Generator/docx2txt.py"

Nothing depends on this for correctness — it exists so that a reviewer, or Claude in
either environment, can see what a rebuild actually altered.
"""
import sys

try:
    from docx import Document
except ImportError:
    print("(python-docx not installed; cannot render this document)")
    sys.exit(0)

d = Document(sys.argv[1])

for section in d.sections:
    for p in section.header.paragraphs:
        if p.text.strip():
            print(f"[header] {p.text.strip()}")

for block in d.element.body.iter():
    tag = block.tag.split("}")[-1]
    if tag == "p":
        # paragraphs inside tables are emitted with the table, not here
        if block.getparent().tag.split("}")[-1] == "tc":
            continue
        text = "".join(n.text or "" for n in block.iter() if n.tag.endswith("}t"))
        if text.strip():
            print(text.strip())
    elif tag == "tbl":
        for row in block.findall(".//{*}tr"):
            cells = []
            for tc in row.findall("./{*}tc"):
                cells.append(" ".join(
                    "".join(n.text or "" for n in p.iter() if n.tag.endswith("}t")).strip()
                    for p in tc.findall(".//{*}p")
                ).strip())
            print(" | ".join(cells))

for section in d.sections:
    for p in section.footer.paragraphs:
        if p.text.strip():
            print(f"[footer] {p.text.strip()}")
